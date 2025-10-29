#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Word检测报告自动提取程序
从Word文档中提取关键信息并输出到Excel文件
"""

import os
import re
import pandas as pd
from datetime import datetime
from docx import Document
from openpyxl import Workbook
from openpyxl.utils.dataframe import dataframe_to_rows


class ReportProcessor:
    """检测报告处理器"""
    
    def __init__(self, input_folder="reports_to_process", output_file="检测结果汇总表_最终版.xlsx"):
        """
        初始化处理器
        
        Args:
            input_folder (str): 输入文件夹路径
            output_file (str): 输出Excel文件名
        """
        self.input_folder = input_folder
        self.output_file = output_file
        self.results = []
        
    def extract_report_number(self, doc_text):
        """
        提取报告编号
        
        Args:
            doc_text (str): 文档文本内容
            
        Returns:
            str: 报告编号，如果未找到返回空字符串
        """
        # 匹配"报告编号："后的内容
        patterns = [
            r'报告编号[：:]\s*([A-Za-z0-9\-]+)',
            r'受控编号[：:]\s*([A-Za-z0-9\-]+)',
            r'编号[：:]\s*([A-Za-z0-9\-]+)'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, doc_text)
            if match:
                return match.group(1).strip()
        
        return ""
    
    def extract_sample_name(self, doc_text):
        """
        提取样品名称
        
        Args:
            doc_text (str): 文档文本内容
            
        Returns:
            str: 样品名称，如果未找到返回空字符串
        """
        # 匹配"样品名称："后的内容
        patterns = [
            r'样品名称[：:]\s*([^\n\r]+)',
            r'工程名称[：:]\s*([^\n\r]+)',
            r'项目名称[：:]\s*([^\n\r]+)'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, doc_text)
            if match:
                return match.group(1).strip()
        
        return ""
    
    def extract_test_date(self, doc_text):
        """
        提取检测日期
        
        Args:
            doc_text (str): 文档文本内容
            
        Returns:
            str: 格式化的检测日期(YYYY-MM-DD)，如果未找到返回"未知"
        """
        # 匹配各种日期格式
        patterns = [
            r'检测日期[：:]\s*(\d{4}年\d{1,2}月\d{1,2}日)',
            r'检测日期[：:]\s*(\d{4}-\d{1,2}-\d{1,2})',
            r'日期[：:]\s*(\d{4}年\d{1,2}月\d{1,2}日)',
            r'Date[：:]\s*(\d{4}-\d{1,2}-\d{1,2})',
            r'(\d{4}年\d{1,2}月\d{1,2}日)',
            r'(\d{4}-\d{1,2}-\d{1,2})'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, doc_text)
            if match:
                date_str = match.group(1)
                # 统一格式化为YYYY-MM-DD
                if '年' in date_str:
                    # 处理"2024年11月18日"格式
                    date_match = re.match(r'(\d{4})年(\d{1,2})月(\d{1,2})日', date_str)
                    if date_match:
                        year, month, day = date_match.groups()
                        return f"{year}-{month.zfill(2)}-{day.zfill(2)}"
                else:
                    # 处理"2024-11-18"格式
                    parts = date_str.split('-')
                    if len(parts) == 3:
                        return f"{parts[0]}-{parts[1].zfill(2)}-{parts[2].zfill(2)}"
        
        return "未知"
    
    def extract_table_data(self, doc):
        """
        从文档表格中提取检测项目数据
        
        Args:
            doc: Document对象
            
        Returns:
            list: 包含检测项目、实测值、单项判定的字典列表
        """
        table_data = []
        
        # 遍历所有表格
        for table in doc.tables:
            # 检查表格是否包含检测相关信息
            has_test_data = False
            header_mapping = {}
            
            # 分析表头 - 检查多行表头
            header_rows = []
            if len(table.rows) > 1:
                # 收集前几行作为可能的表头
                for i in range(min(3, len(table.rows))):
                    header_rows.append(table.rows[i])
                
                # 分析表头
                for row_idx, header_row in enumerate(header_rows):
                    for i, cell in enumerate(header_row.cells):
                        cell_text = cell.text.strip()
                        # 智能匹配表头
                        if any(keyword in cell_text for keyword in ['检测项目', '项目', '检测内容', '检测部位', '部位']):
                            header_mapping['test_item'] = i
                            has_test_data = True
                        elif any(keyword in cell_text for keyword in ['实测值', '实测结果', '数值', '结果', '强度', '芯样抗压强度', '代表值']):
                            header_mapping['measured_value'] = i
                            has_test_data = True
                        elif any(keyword in cell_text for keyword in ['单项判定', '单项结论', '判定', '结论', '结论']):
                            header_mapping['judgment'] = i
                            has_test_data = True
                
                # 如果找到相关表头，提取数据行
                if has_test_data and len(header_mapping) >= 2:
                    # 从表头下一行开始提取数据
                    start_row = len(header_rows)
                    current_location = ""  # 当前检测部位
                    location_groups = {}  # 按检测部位分组的数据
                    
                    for row_idx, row in enumerate(table.rows[start_row:]):  # 跳过表头
                        if len(row.cells) > max(header_mapping.values()):
                            item_data = {}
                            
                            # 提取检测项目
                            if 'test_item' in header_mapping:
                                test_item_text = row.cells[header_mapping['test_item']].text.strip()
                                # 合并多行单元格的情况
                                if not test_item_text and row_idx > 0:
                                    # 如果当前行为空，尝试使用上一行的检测部位
                                    test_item_text = current_location
                                else:
                                    current_location = test_item_text
                                item_data['test_item'] = test_item_text
                            
                            # 提取实测值
                            if 'measured_value' in header_mapping:
                                measured_value = row.cells[header_mapping['measured_value']].text.strip()
                                # 如果是代表值列，优先使用代表值
                                item_data['measured_value'] = measured_value
                            
                            # 提取单项判定
                            if 'judgment' in header_mapping:
                                judgment_text = row.cells[header_mapping['judgment']].text.strip()
                                # 合并多行单元格的情况
                                if not judgment_text and row_idx > 0:
                                    # 如果当前行为空，尝试使用上一行的判定结果
                                    judgment_text = table_data[-1]['judgment'] if table_data else ""
                                item_data['judgment'] = judgment_text
                            
                            # 过滤无效数据
                            test_item = item_data.get('test_item', '')
                            measured_value = item_data.get('measured_value', '')
                            judgment = item_data.get('judgment', '')
                            
                            # 排除空白行和"以下空白"等无效内容
                            if (test_item and 
                                test_item not in ['以下空白', '', ' '] and
                                not test_item.startswith('以下') and
                                measured_value and
                                measured_value not in ['', ' '] and
                                judgment and
                                judgment not in ['', ' ']):
                                
                                # 按检测部位分组，每个部位保留所有不同的实测值
                                location_key = test_item
                                if location_key not in location_groups:
                                    location_groups[location_key] = []
                                location_groups[location_key].append(item_data)
                    
                    # 处理分组数据，为每个检测部位创建记录
                    for location, items in location_groups.items():
                        if items:
                            # 提取所有不同的代表值，按实测值分组
                            representative_values = {}
                            
                            for item in items:
                                val_text = item['measured_value']
                                judgment = item['judgment']
                                
                                # 提取数值，四舍五入到整数用于分组
                                val_match = re.search(r'(\d+\.?\d*)', val_text)
                                if val_match:
                                    numeric_value = float(val_match.group(1))
                                    # 使用数值的四舍五入值作为分组依据
                                    rounded_value = round(numeric_value)
                                    key = f"{location}_{rounded_value}"
                                    if key not in representative_values:
                                        representative_values[key] = {
                                            'location': location,
                                            'judgment': judgment,
                                            'values': [],
                                            'sample_item': item
                                        }
                                    representative_values[key]['values'].append(numeric_value)
                            
                            # 为每个代表值创建记录
                            for key, data in representative_values.items():
                                if data['values']:
                                    # 计算该组代表值的平均值
                                    avg_value = sum(data['values']) / len(data['values'])
                                    
                                    # 创建记录
                                    representative_item = data['sample_item'].copy()
                                    # 保留原始格式但更新数值
                                    original_format = data['sample_item']['measured_value']
                                    if re.match(r'\d+\.?\d*', original_format):
                                        representative_item['measured_value'] = f"{avg_value:.1f}"
                                    else:
                                        # 如果原始格式包含单位，保持单位
                                        unit_match = re.search(r'([^\d]+)$', original_format)
                                        if unit_match:
                                            representative_item['measured_value'] = f"{avg_value:.1f} {unit_match.group(1).strip()}"
                                        else:
                                            representative_item['measured_value'] = f"{avg_value:.1f}"
                                    
                                    table_data.append(representative_item)
            
            # 如果找到包含检测数据的表格，停止搜索
            if table_data:
                break
        
        # 如果没有找到表格数据，尝试从文本中提取
        if not table_data:
            doc_text = "\n".join([para.text for para in doc.paragraphs])
            table_data = self.extract_from_text(doc_text)
        
        return table_data
    
    def extract_from_text(self, doc_text):
        """
        从文本中提取检测数据（当没有表格时使用）
        
        Args:
            doc_text (str): 文档文本内容
            
        Returns:
            list: 包含检测项目、实测值、单项判定的字典列表
        """
        text_data = []
        
        # 匹配"项目名：数值（判定）"格式
        pattern = r'([^\n:：]+)[：:]\s*([^\n（\(]+)[（\(]([^\n）\)]+)[）\)]'
        matches = re.findall(pattern, doc_text)
        
        for match in matches:
            test_item = match[0].strip()
            measured_value = match[1].strip()
            judgment = match[2].strip()
            
            # 过滤无效数据
            if (test_item and 
                test_item not in ['以下空白', '', ' '] and
                not test_item.startswith('以下') and
                measured_value and
                measured_value not in ['', ' '] and
                judgment and
                judgment not in ['', ' ']):
                
                item_data = {
                    'test_item': test_item,
                    'measured_value': measured_value,
                    'judgment': judgment
                }
                text_data.append(item_data)
        
        # 尝试匹配混凝土强度检测结果格式
        # 例如："芯样抗压强度(MPa)"和"结论"列
        concrete_pattern = r'芯样抗压强度.*?结论.*?(\d+\.\d+)\s*MPa.*?(合格|不合格)'
        concrete_matches = re.findall(concrete_pattern, doc_text, re.DOTALL)
        
        for match in concrete_matches:
            item_data = {
                'test_item': '混凝土抗压强度',
                'measured_value': match[0] + ' MPa',
                'judgment': match[1]
            }
            text_data.append(item_data)
        
        # 尝试从附表中提取数据
        # 查找"附表"或表格数据
        table_pattern = r'序号.*?检测部位.*?强度等级.*?芯样抗压强度.*?结论.*?(\d+).*?([^\s]+).*?C\d+.*?(\d+\.\d+)\s*MPa.*?(合格|不合格)'
        table_matches = re.findall(table_pattern, doc_text, re.DOTALL)
        
        for match in table_matches:
            item_data = {
                'test_item': f'混凝土抗压强度(部位{match[0]})',
                'measured_value': match[2] + ' MPa',
                'judgment': match[3]
            }
            text_data.append(item_data)
        
        return text_data
    
    def process_document(self, file_path):
        """
        处理单个Word文档
        
        Args:
            file_path (str): 文档文件路径
            
        Returns:
            dict: 提取的数据，如果处理失败返回None
        """
        try:
            # 读取Word文档
            doc = Document(file_path)
            
            # 获取文档文本
            doc_text = "\n".join([para.text for para in doc.paragraphs])
            
            # 提取基本信息
            report_number = self.extract_report_number(doc_text)
            sample_name = self.extract_sample_name(doc_text)
            test_date = self.extract_test_date(doc_text)
            
            # 提取表格数据
            table_data = self.extract_table_data(doc)
            
            # 如果没有找到检测数据，返回None
            if not table_data:
                print(f"警告：文件 {os.path.basename(file_path)} 中未找到检测数据")
                return None
            
            # 为每个检测项目创建一条记录
            results = []
            for item in table_data:
                result = {
                    '报告编号': report_number,
                    '样品名称': sample_name,
                    '检测项目': item.get('test_item', ''),
                    '实测值': item.get('measured_value', ''),
                    '单项判定': item.get('judgment', ''),
                    '检测日期': test_date
                }
                results.append(result)
            
            return results
            
        except Exception as e:
            print(f"处理文件 {os.path.basename(file_path)} 时出错: {str(e)}")
            return None
    
    def process_all_documents(self):
        """
        处理文件夹中的所有Word文档
        
        Returns:
            list: 所有提取的结果
        """
        if not os.path.exists(self.input_folder):
            print(f"错误：文件夹 {self.input_folder} 不存在")
            return []
        
        docx_files = [f for f in os.listdir(self.input_folder) if f.endswith('.docx') and not f.startswith('~$')]
        
        if not docx_files:
            print(f"错误：文件夹 {self.input_folder} 中没有找到.docx文件")
            return []
        
        print(f"找到 {len(docx_files)} 个Word文档，开始处理...")
        
        for filename in docx_files:
            file_path = os.path.join(self.input_folder, filename)
            print(f"正在处理: {filename}")
            
            results = self.process_document(file_path)
            if results:
                self.results.extend(results)
                print(f"  成功提取 {len(results)} 条记录")
            else:
                print(f"  处理失败或未找到有效数据")
        
        print(f"处理完成，共提取 {len(self.results)} 条记录")
        return self.results
    
    def save_to_excel(self):
        """
        将结果保存到Excel文件
        """
        if not self.results:
            print("没有数据可保存")
            return
        
        try:
            # 创建DataFrame
            df = pd.DataFrame(self.results)
            
            # 创建Excel工作簿
            wb = Workbook()
            ws = wb.active
            ws.title = "检测结果汇总"
            
            # 写入数据
            for r in dataframe_to_rows(df, index=False, header=True):
                ws.append(r)
            
            # 调整列宽
            for column in ws.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                ws.column_dimensions[column_letter].width = adjusted_width
            
            # 保存文件
            wb.save(self.output_file)
            print(f"结果已保存到: {self.output_file}")
            
        except Exception as e:
            print(f"保存Excel文件时出错: {str(e)}")


def main():
    """主函数"""
    print("Word检测报告自动提取程序")
    print("=" * 50)
    
    # 创建处理器
    processor = ReportProcessor()
    
    # 处理所有文档
    processor.process_all_documents()
    
    # 保存结果
    processor.save_to_excel()
    
    print("=" * 50)
    print("程序执行完成")


if __name__ == "__main__":
    main()